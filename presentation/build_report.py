from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

NAVY = RGBColor(0x00, 0x35, 0x80)
INK = RGBColor(0x0A, 0x0F, 0x1A)

doc = Document()

normal = doc.styles["Normal"]
normal.font.name = "Calibri"
normal.font.size = Pt(11)
normal.font.color.rgb = INK

for section in doc.sections:
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin = Cm(3)
    section.right_margin = Cm(2.5)


def centered(text, size=11, bold=False, color=INK, space_after=6, font=None):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.line_spacing = 1.15
    r = p.add_run(text)
    r.font.size = Pt(size)
    r.font.bold = bold
    r.font.color.rgb = color
    if font:
        r.font.name = font
    return p


def body(text):
    p = doc.add_paragraph(text)
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.first_line_indent = Cm(1.25)
    p.paragraph_format.line_spacing = 1.5
    p.paragraph_format.space_after = Pt(6)
    return p


def bullet(text):
    p = doc.add_paragraph(text, style="List Bullet")
    p.paragraph_format.line_spacing = 1.5
    p.paragraph_format.space_after = Pt(4)
    return p


def chap(text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(18)
    p.paragraph_format.space_after = Pt(12)
    r = p.add_run(text)
    r.font.size = Pt(14)
    r.font.bold = True
    r.font.color.rgb = NAVY
    return p


def sub(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(6)
    r = p.add_run(text)
    r.font.size = Pt(12)
    r.font.bold = True
    r.font.color.rgb = INK
    return p


def page_break():
    doc.add_page_break()


# ============================== TITLE PAGE
for _ in range(3):
    centered("", space_after=12)
centered("LAPORAN MAGANG", size=16, bold=True, space_after=18)
centered("Modernisasi Portal CSIRT DKI Jakarta:", size=15, bold=True, space_after=6)
centered("Redesain UI/UX, Penerapan Aksesibilitas WCAG 2.1,", size=15, bold=True, space_after=6)
centered("dan Migrasi Framework Yii ke Laravel", size=15, bold=True, space_after=24)
centered("Disusun oleh:", size=12, space_after=6)
centered("Abdul Latif", size=14, bold=True, space_after=2)
centered("NIM: [NIM BINUS University]", size=12, space_after=18)
centered("Program Studi Computer Science", size=12, space_after=2)
centered("BINUS University", size=12, space_after=24)
centered("Instansi: Diskominfotik DKI Jakarta", size=12, space_after=2)
centered("(CSIRT DKI Jakarta — csirt.jakarta.go.id)", size=12, space_after=6)
centered("Periode Magang: Februari – Agustus 2026", size=12, space_after=2)
centered("Pembimbing Lapangan: Pak Andy & Bu Rina", size=12, space_after=24)
centered("2026", size=12, bold=True)
page_break()

# ============================== APPROVAL PAGE
centered("LEMBAR PENGESAHAN", size=14, bold=True, space_after=18)
body(
    "Laporan magang berjudul \u201cModernisasi Portal CSIRT DKI Jakarta: Redesain UI/UX, "
    "Penerapan Aksesibilitas WCAG 2.1, dan Migrasi Framework Yii ke Laravel\u201d ini disusun "
    "oleh Abdul Latif, mahasiswa Program Studi Computer Science BINUS University, sebagai "
    "pertanggungjawaban pelaksanaan magang di Diskominfotik DKI Jakarta pada periode "
    "Februari – Agustus 2026. Laporan ini telah disetujui oleh pembimbing lapangan."
)
doc.add_paragraph()
approval = doc.add_table(rows=1, cols=2)
approval.autofit = True
left = approval.cell(0, 0).paragraphs[0]
left.alignment = WD_ALIGN_PARAGRAPH.LEFT
left.paragraph_format.line_spacing = 1.5
left.add_run("Mengetahui,\nPembimbing Lapangan,\n\n\n\n( Pak Andy / Bu Rina )")
right = approval.cell(0, 1).paragraphs[0]
right.alignment = WD_ALIGN_PARAGRAPH.RIGHT
right.paragraph_format.line_spacing = 1.5
right.add_run("Jakarta, Agustus 2026\nMahasiswa,\n\n\n\n( Abdul Latif )")
page_break()

# ============================== PREFACE
centered("KATA PENGANTAR", size=14, bold=True, space_after=18)
body(
    "Puji syukur penulis panjatkan ke hadirat Tuhan Yang Maha Esa atas rahmat dan karunia-Nya "
    "sehingga laporan magang ini dapat diselesaikan. Laporan ini disusun sebagai bentuk "
    "pertanggungjawaban pelaksanaan magang di Diskominfotik DKI Jakarta, khususnya pada satuan "
    "tugas CSIRT DKI Jakarta, selama periode Februari – Agustus 2026."
)
body(
    "Selama pelaksanaan magang, penulis terlibat dalam modernisasi portal CSIRT DKI Jakarta, "
    "meliputi migrasi framework dari Yii ke Laravel, perancangan ulang antarmuka dan pengalaman "
    "pengguna, penerapan aksesibilitas sesuai standar WCAG 2.1, serta penguatan keamanan dan "
    "stabilitas aplikasi. Pengalaman tersebut memberikan pemahaman mendalam mengenai proses "
    "pengembangan sistem informasi di lingkungan instansi pemerintah."
)
body(
    "Penulis mengucapkan terima kasih kepada Bapak Andy dan Ibu Rina selaku pembimbing lapangan "
    "yang telah memberikan arahan, masukan, dan dukungan selama pelaksanaan magang, kepada jajaran "
    "Diskominfotik DKI Jakarta atas kesempatan yang diberikan, serta kepada seluruh pihak yang "
    "turut membantu penyelesaian laporan ini. Penulis menyadari laporan ini masih memiliki "
    "kekurangan, sehingga kritik dan saran yang membangun sangat diharapkan."
)
doc.add_paragraph()
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
p.paragraph_format.line_spacing = 1.5
p.add_run("Jakarta, Agustus 2026\nPenulis,\n\n\n( Abdul Latif )")
page_break()

# ============================== TOC
centered("DAFTAR ISI", size=14, bold=True, space_after=18)
toc = [
    "KATA PENGANTAR",
    "DAFTAR ISI",
    "BAB 1 PENDAHULUAN",
    "1.1 Latar Belakang",
    "1.2 Tujuan Magang",
    "1.3 Ruang Lingkup",
    "1.4 Sistematika Penulisan",
    "BAB 2 PROFIL INSTANSI",
    "2.1 Profil Diskominfotik DKI Jakarta",
    "2.2 Profil CSIRT DKI Jakarta",
    "2.3 Tugas dan Layanan CSIRT DKI Jakarta",
    "BAB 3 ANALISIS SISTEM YANG ADA",
    "3.1 Kondisi Sistem Sebelumnya",
    "3.2 Identifikasi Permasalahan",
    "3.3 Analisis Kebutuhan Pengguna",
    "BAB 4 METODOLOGI DAN PELAKSANAAN MAGANG",
    "4.1 Metodologi Pelaksanaan",
    "4.2 Alat dan Teknologi",
    "4.3 Jadwal Pelaksanaan",
    "4.4 Dokumentasi Pelaksanaan",
    "BAB 5 HASIL DAN PEMBAHASAN",
    "5.1 Perancangan Design System",
    "5.2 Redesain Antarmuka Publik",
    "5.3 Alur Pelaporan Insiden (Multi-Step Wizard)",
    "5.4 Penerapan Aksesibilitas WCAG 2.1",
    "5.5 Penguatan Keamanan dan Stabilitas",
    "5.6 Fitur Pendukung",
    "5.7 Pembahasan Umum",
    "BAB 6 PENUTUP",
    "6.1 Kesimpulan",
    "6.2 Saran",
]
for item in toc:
    p = doc.add_paragraph(item)
    p.paragraph_format.line_spacing = 1.4
    p.paragraph_format.space_after = Pt(2)
    if item.startswith("BAB"):
        p.runs[0].font.bold = True
page_break()

# ============================== BAB 1
chap("BAB 1 PENDAHULUAN")
sub("1.1 Latar Belakang")
body(
    "Perkembangan ancaman keamanan siber yang semakin masif menuntut setiap instansi pemerintah "
    "untuk memperkuat kesiapan dalam menghadapi insiden siber. Dalam konteks tersebut, Dinas "
    "Komunikasi, Informatika, dan Statistik (Diskominfotik) DKI Jakarta melalui Computer Security "
    "Incident Response Team (CSIRT) DKI Jakarta menyediakan portal publik sebagai kanal utama bagi "
    "masyarakat dan pegawai pemerintah untuk melaporkan insiden siber, memperoleh berita keamanan, "
    "serta memanfaatkan berbagai sumber daya literasi keamanan digital."
)
body(
    "Pada kondisi awal, portal CSIRT DKI Jakarta dibangun menggunakan framework Yii dengan "
    "antarmuka yang datar dan kurang modern. Alur pelaporan insiden disajikan dalam satu formulir "
    "panjang yang rumit, sehingga menyulitkan pengguna non-teknis. Selain itu, tingkat aksesibilitas "
    "portal masih rendah dan belum sepenuhnya memenuhi standar Web Content Accessibility Guidelines "
    "(WCAG) 2.1, sehingga kurang inklusif, khususnya bagi penyandang disabilitas."
)
body(
    "Berdasarkan kondisi tersebut, pelaksanaan magang difokuskan pada modernisasi portal CSIRT DKI "
    "Jakarta melalui tiga hal utama: migrasi framework dari Yii ke Laravel, perancangan ulang "
    "antarmuka dan pengalaman pengguna (UI/UX), serta penerapan aksesibilitas sesuai standar "
    "WCAG 2.1."
)
sub("1.2 Tujuan Magang")
bullet("Melakukan migrasi codebase portal dari framework Yii ke Laravel 12 agar lebih mudah dikembangkan dan dipelihara.")
bullet("Merancang ulang antarmuka dan pengalaman pengguna (UI/UX) portal berdasarkan design system yang konsisten.")
bullet("Menyederhanakan alur pelaporan insiden menjadi wizard multi-langkah yang intuitif bagi pengguna non-teknis.")
bullet("Menerapkan prinsip aksesibilitas WCAG 2.1, termasuk mode kontras tinggi dan mode gelap.")
bullet("Memperkuat keamanan dan stabilitas aplikasi, meliputi validasi, penanganan error, dan rate limiting.")
sub("1.3 Ruang Lingkup")
body(
    "Ruang lingkup magang mencakup perancangan design system, redesain seluruh halaman publik "
    "(beranda, berita, kegiatan, peringatan keamanan, infografis, peraturan, panduan teknis, profil, "
    "dan pencarian), perancangan alur pelaporan insiden, pengembangan panel admin, penerapan "
    "aksesibilitas, penguatan keamanan formulir, serta penyusunan dokumentasi teknis proyek."
)
sub("1.4 Sistematika Penulisan")
body(
    "Laporan ini disusun dalam enam bab. Bab 1 memaparkan pendahuluan yang mencakup latar belakang, "
    "tujuan, ruang lingkup, dan sistematika penulisan. Bab 2 membahas profil instansi tempat magang. "
    "Bab 3 menganalisis sistem yang ada sebelum dilakukan modernisasi. Bab 4 menguraikan metodologi "
    "dan pelaksanaan magang. Bab 5 menyajikan hasil pekerjaan beserta pembahasannya. Bab 6 memuat "
    "kesimpulan dan saran."
)
page_break()

# ============================== BAB 2
chap("BAB 2 PROFIL INSTANSI")
sub("2.1 Profil Diskominfotik DKI Jakarta")
body(
    "Diskominfotik DKI Jakarta adalah perangkat daerah Pemerintah Provinsi DKI Jakarta yang "
    "melaksanakan urusan pemerintahan bidang komunikasi, informatika, dan statistik. Instansi ini "
    "mengelola layanan digital serta infrastruktur teknologi informasi bagi seluruh organisasi "
    "perangkat daerah di lingkungan Pemprov DKI Jakarta, termasuk pengamanan data dan sistem "
    "elektronik, pengelolaan informasi publik, serta dukungan statistik bagi pengambilan keputusan."
)
sub("2.2 Profil CSIRT DKI Jakarta")
body(
    "CSIRT DKI Jakarta merupakan tim tanggap insiden keamanan komputer di lingkungan Pemerintah "
    "Provinsi DKI Jakarta. Tim ini berperan sebagai garda terdepan dalam merespons insiden siber, "
    "dengan tugas menerima dan menganalisis laporan insiden, menerbitkan peringatan dan berita "
    "keamanan, menyelenggarakan edukasi literasi keamanan digital, serta mendukung kebijakan "
    "keamanan informasi di lingkungan instansi pemerintah. Portal csirt.jakarta.go.id menjadi "
    "antarmuka utama antara tim CSIRT dengan masyarakat dan pegawai pemerintah."
)
sub("2.3 Tugas dan Layanan CSIRT DKI Jakarta")
bullet("Menerima dan menindaklanjuti laporan insiden siber dari instansi pemerintah dan masyarakat.")
bullet("Menyebarluaskan peringatan keamanan (advisory) dan berita keamanan siber terkini.")
bullet("Menyediakan materi edukasi berupa infografis, panduan teknis, serta peraturan dan kebijakan terkait.")
bullet("Melayani konsultasi keamanan informasi melalui kanal resmi, termasuk hotline 24/7.")
page_break()

# ============================== BAB 3
chap("BAB 3 ANALISIS SISTEM YANG ADA")
sub("3.1 Kondisi Sistem Sebelumnya")
body(
    "Sebelum dilakukan modernisasi, portal CSIRT DKI Jakarta dibangun menggunakan framework Yii "
    "dengan arsitektur monolitik. Konten dikelola melalui panel admin sederhana, sedangkan halaman "
    "publik menampilkan daftar berita, peringatan keamanan, kegiatan, infografis, peraturan, dan "
    "panduan teknis. Formulir pelaporan insiden dirancang dalam satu halaman dengan banyak kolom "
    "yang ditampilkan sekaligus."
)
sub("3.2 Identifikasi Permasalahan")
bullet("Antarmuka datar dan kurang modern sehingga menurunkan kepercayaan dan kenyamanan pengguna.")
bullet("Formulir pelaporan insiden yang panjang dan rumit, tidak ramah bagi pengguna non-teknis.")
bullet("Tidak terdapat mode aksesibilitas seperti kontras tinggi dan mode gelap.")
bullet("Penanganan error dan validasi pada formulir belum optimal.")
bullet("Belum ada pembatasan laju permintaan (rate limiting) pada formulir publik.")
bullet("Dokumentasi teknis terbatas sehingga menyulitkan pengembangan lanjutan.")
sub("3.3 Analisis Kebutuhan Pengguna")
body(
    "Pengguna publik, yaitu masyarakat dan pegawai pemerintah, membutuhkan kanal pelaporan insiden "
    "yang sederhana dan cepat, serta informasi keamanan yang mudah diakses, termasuk oleh penyandang "
    "disabilitas. Pengguna admin, yaitu petugas CSIRT, membutuhkan panel pengelolaan konten yang "
    "efisien dan mudah dinavigasikan. Selain itu, pengembang membutuhkan basis kode yang modern "
    "serta dokumentasi yang memadai agar proyek berkelanjutan."
)
page_break()

# ============================== BAB 4
chap("BAB 4 METODOLOGI DAN PELAKSANAAN MAGANG")
sub("4.1 Metodologi Pelaksanaan")
body("Pelaksanaan magang dilakukan melalui lima tahap utama dengan pendekatan iteratif:")
bullet("Analisis: studi terhadap sistem eksisting berbasis Yii, identifikasi kebutuhan pengguna, dan perumusan lingkup pekerjaan.")
bullet("Perancangan: penyusunan design system berbasis token, perancangan alur wizard pelaporan insiden, dan wireframe antarmuka.")
bullet("Implementasi: migrasi dan pengembangan menggunakan Laravel 12, Blade, serta CSS custom properties.")
bullet("Pengujian: pengujian fungsional, pemeriksaan kontras aksesibilitas, dan pengujian responsivitas lintas perangkat.")
bullet("Dokumentasi: penulisan dokumentasi proyek secara paralel agar konteks pengembangan tetap terjaga.")
sub("4.2 Alat dan Teknologi")
bullet("Laravel 12 dengan PHP 8.2 sebagai framework utama hasil migrasi.")
bullet("Blade templating dan Bootstrap 5.3 untuk struktur antarmuka.")
bullet("CSS custom properties (design tokens) untuk pengaturan warna, tipografi, dan jarak.")
bullet("SQLite sebagai basis data pengembangan.")
bullet("Git dan GitHub untuk kontrol versi dan kolaborasi.")
bullet("Figma dan Canva untuk perancangan antarmuka.")
sub("4.3 Jadwal Pelaksanaan")
jadwal = [
    ("Periode", "Kegiatan"),
    ("Februari – Maret", "Studi sistem eksisting (codebase Yii), persiapan lingkungan pengembangan Laravel, dan perumusan kebutuhan"),
    ("April – Mei", "Perancangan design system, wireframe antarmuka, dan alur wizard pelaporan insiden"),
    ("Juni", "Implementasi redesain antarmuka publik dan komponen halaman"),
    ("Juli", "Implementasi aksesibilitas dan wizard pelaporan; pengujian serta perbaikan; penyusunan testimoni magang (27 Juli)"),
    ("Agustus", "Finalisasi dokumentasi, penyempurnaan aplikasi, penyusunan laporan magang, dan proposal pre-thesis (22 Agustus)"),
]
table = doc.add_table(rows=len(jadwal), cols=2)
table.style = "Table Grid"
for i, (a, b) in enumerate(jadwal):
    cell_a, cell_b = table.cell(i, 0), table.cell(i, 1)
    cell_a.paragraphs[0].add_run(a).font.bold = (i == 0)
    cell_b.paragraphs[0].add_run(b)
    cell_a.paragraphs[0].paragraph_format.line_spacing = 1.2
    cell_b.paragraphs[0].paragraph_format.line_spacing = 1.2
    for cell in (cell_a, cell_b):
        for para in cell.paragraphs:
            para.paragraph_format.space_after = Pt(2)
doc.add_paragraph()
sub("4.4 Dokumentasi Pelaksanaan")
body(
    "Selama pelaksanaan, seluruh konvensi dan konteks proyek didokumentasikan di dalam repositori, "
    "antara lain AGENTS.md (konvensi dan perintah proyek), FEATURES.md (inventarisasi fitur), "
    "DESIGN_SYSTEM.md (sistem desain), SETUP_GUIDE.md (panduan instalasi), serta "
    "ACCESSIBILITY_VERIFICATION.md (verifikasi aksesibilitas), sehingga memudahkan pengembang "
    "berikutnya dalam melanjutkan pekerjaan."
)
page_break()

# ============================== BAB 5
chap("BAB 5 HASIL DAN PEMBAHASAN")
sub("5.1 Perancangan Design System")
body(
    "Sebagai fondasi seluruh pekerjaan desain, disusun design system yang mendefinisikan token "
    "warna, tipografi, dan jarak dalam bentuk CSS custom properties. Palet utama menggunakan biru "
    "DKI Jakarta (navy #003580) sebagai warna primer, near-black #0A0F1A untuk teks, serta abu "
    "muda #F4F5F7 untuk latar bagian. Tipografi menggunakan Plus Jakarta Sans untuk judul dan "
    "Inter untuk teks isi. Seluruh halaman wajib menggunakan token ini sehingga konsistensi "
    "visual terjaga dan dapat diubah secara terpusat, termasuk untuk mode aksesibilitas."
)
sub("5.2 Redesain Antarmuka Publik")
body(
    "Seluruh halaman publik didesain ulang dengan pola header gelap yang konsisten, hero ringkas "
    "pada beranda dengan ajakan untuk melapor insiden, pengumuman peringatan keamanan aktif, "
    "carousel berita terbaru, serta grid layanan yang menghubungkan pengguna ke peringatan, "
    "infografis, peraturan, dan panduan teknis. Sistem kartu kustom (bukan kartu Bootstrap) "
    "digunakan untuk konsistensi tampilan, sedangkan Bootstrap hanya dipakai untuk grid, formulir, "
    "dan tabel."
)
sub("5.3 Alur Pelaporan Insiden (Multi-Step Wizard)")
body(
    "Formulir pelaporan insiden diubah dari satu halaman panjang menjadi wizard tiga langkah "
    "tanpa berpindah halaman: langkah pertama mengisi data pelapor (nama, surel, telepon, tanggal "
    "ditemukan), langkah kedua data situs (domain dan URL), serta langkah ketiga detail insiden "
    "(deskripsi, jenis dan tingkat risiko, skor CVSS, bukti, dan rekomendasi). Kolom risiko "
    "dibuat opsional agar ramah pengguna non-teknis, terdapat verifikasi CAPTCHA bertuliskan "
    "\u201cJKT\u201d, serta unggah bukti berupa berkas PNG/JPG maksimal 2 MB. Setelah terkirim, "
    "pengguna diarahkan ke halaman terima kasih."
)
sub("5.4 Penerapan Aksesibilitas WCAG 2.1")
body(
    "Widget aksesibilitas dipasang pada seluruh halaman dengan tiga mode: kontras tinggi, mode "
    "gelap, dan mode default. Implementasinya memanfaatkan CSS custom properties yang ditimpa "
    "melalui berkas khusus, dengan JavaScript yang mengalihkan kelas pada elemen root. Dengan "
    "demikian, seluruh komponen yang menggunakan token akan menyesuaikan warnanya secara "
    "otomatis, memenuhi prinsip kontras warna pada WCAG 2.1."
)
sub("5.5 Penguatan Keamanan dan Stabilitas")
body(
    "Seluruh pengiriman formulir dan operasi CRUD admin dibungkus dalam penanganan error "
    "(try/catch) sehingga kegagalan tidak menampilkan halaman error mentah, melainkan pesan "
    "berbahasa Indonesia yang informatif. Formulir publik dibatasi dengan rate limiting sebanyak "
    "60 permintaan per menit per IP, dan validasi dilakukan di sisi server. Rute admin dilindungi "
    "middleware autentikasi dan pemeriksaan status admin."
)
sub("5.6 Fitur Pendukung")
body(
    "Fitur pendukung yang turut dikerjakan antara lain pencarian site-wide yang mencakup enam "
    "jenis konten (berita, peringatan, kegiatan, infografis, peraturan, panduan) dengan kata kunci "
    "minimal dua karakter, panel admin dengan tab dan pagination 15 baris per tab beserta status "
    "tab yang bertahan antar halaman, serta data seed untuk kebutuhan pengembangan dan pengujian."
)
sub("5.7 Pembahasan Umum")
body(
    "Setelah modernisasi, portal CSIRT DKI Jakarta memiliki antarmuka yang modern, konsisten, dan "
    "mudah diakses oleh beragam kalangan pengguna, termasuk penyandang disabilitas. Alur pelaporan "
    "insiden menjadi lebih sederhana dan mengarahkan pengguna langkah demi langkah. Basis kode yang "
    "berpindah ke Laravel 12 lebih terstruktur dan mudah dipelihara, didukung dokumentasi lengkap "
    "sehingga pekerjaan dapat dilanjutkan oleh pengembang lain. Pekerjaan ini sekaligus menjadi "
    "bahan kajian untuk proposal pre-thesis mengenai aksesibilitas, pengalaman pengguna "
    "multi-langkah, dan migrasi framework."
)
page_break()

# ============================== BAB 6
chap("BAB 6 PENUTUP")
sub("6.1 Kesimpulan")
body(
    "Berdasarkan pelaksanaan magang di Diskominfotik DKI Jakarta, dapat disimpulkan beberapa hal "
    "berikut: (1) migrasi framework dari Yii ke Laravel 12 berhasil dilakukan sehingga portal lebih "
    "mudah dikembangkan dan dipelihara; (2) redesain UI/UX dengan design system berbasis token "
    "menghasilkan antarmuka yang modern, konsisten, dan mudah digunakan; (3) alur pelaporan insiden "
    "multi-langkah mempermudah pengguna non-teknis dalam menyampaikan laporan; (4) penerapan "
    "aksesibilitas WCAG 2.1, termasuk mode kontras tinggi dan mode gelap, membuat portal lebih "
    "inklusif; serta (5) penguatan keamanan melalui validasi, penanganan error, dan rate limiting "
    "meningkatkan stabilitas dan kepercayaan pengguna."
)
sub("6.2 Saran")
body(
    "Untuk pengembangan selanjutnya, disarankan agar instansi melakukan pengujian pengguna (user "
    "testing) dengan perwakilan masyarakat dan penyandang disabilitas, memberikan pelatihan "
    "pengelolaan konten bagi petugas admin, serta menyusun mekanisme pencadangan data. Pengembangan "
    "lanjutan dapat mencakup notifikasi otomatis status penanganan laporan, integrasi dengan kanal "
    "media sosial, dan penguatan keamanan infrastruktur."
)

OUT = r"K:\GitHub\new-csirt\presentation\Laporan_Magang_CSIRT_DKI_Jakarta.docx"
doc.save(OUT)
print("saved:", OUT)
